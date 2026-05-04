import io
import re
import pdfplumber
import docx
from bs4 import BeautifulSoup

def _clean_text(text: str) -> str:
    # Remove non-printable characters
    text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]', '', text)
    # Collapse multiple spaces
    text = re.sub(r'[ \t]+', ' ', text)
    # Collapse 3+ newlines to 2
    text = re.sub(r'\n{3,}', '\n\n', text)
    # Strip each line
    lines = [line.strip() for line in text.split('\n')]
    return '\n'.join(line for line in lines if line)

def _extract_pdf(file_bytes: bytes) -> str:
    extracted_text = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            words = page.extract_words(
                x_tolerance=3,
                y_tolerance=3,
                keep_blank_chars=False,
                use_text_flow=True
            )
            # Sort words by top (y) then left (x)
            words.sort(key=lambda w: (round(w['top'], 1), round(w['x0'], 1)))
            
            # Reconstruct text line by line
            current_line = []
            last_top = None
            
            for word in words:
                top = round(word['top'], 1)
                if last_top is None:
                    last_top = top
                
                if abs(top - last_top) > 5:  # New line threshold
                    extracted_text.append(' '.join(current_line))
                    current_line = []
                    last_top = top
                
                current_line.append(word['text'])
            
            if current_line:
                extracted_text.append(' '.join(current_line))
            
            extracted_text.append('\n') # Page break separator
            
    return '\n'.join(extracted_text)

def _extract_docx(file_bytes: bytes) -> str:
    doc = docx.Document(io.BytesIO(file_bytes))
    extracted_text = []
    
    # Extract paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            extracted_text.append(para.text.strip())
            
    # Extract tables
    for table in doc.tables:
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_data:
                extracted_text.append(" | ".join(row_data))
                
    return '\n'.join(extracted_text)

def _extract_html(file_bytes: bytes) -> str:
    soup = BeautifulSoup(file_bytes.decode('utf-8', errors='replace'), 'lxml')
    for tag in soup(['script', 'style', 'meta', 'link', 'head']):
        tag.decompose()
    
    # Extract text with newline separator for blocks
    text = soup.get_text(separator='\n')
    return text

def extract_text(file_bytes: bytes, filename: str) -> tuple[str, str]:
    ext = filename.split('.')[-1].lower()
    
    if ext == 'pdf':
        raw_text = _extract_pdf(file_bytes)
        format_type = 'pdf'
    elif ext in ['docx', 'doc']:
        raw_text = _extract_docx(file_bytes)
        format_type = 'docx'
    elif ext in ['html', 'htm']:
        raw_text = _extract_html(file_bytes)
        format_type = 'html'
    else:
        raise ValueError(f"Unsupported file format: {ext}")
        
    cleaned_text = _clean_text(raw_text)
    return cleaned_text, format_type

def segment_sections(text: str) -> dict[str, str]:
    # Common regex patterns for sections
    section_patterns = {
        "summary": r"(?i)^(summary|profile|objective|about me|professional summary)",
        "experience": r"(?i)^(experience|work experience|employment|career history|professional experience)",
        "education": r"(?i)^(education|academic|qualifications|degrees)",
        "skills": r"(?i)^(skills|technical skills|competencies|core competencies|expertise)",
        "projects": r"(?i)^(projects|key projects|selected projects|portfolio)",
        "certifications": r"(?i)^(certifications|certificates|licenses|credentials)",
        "achievements": r"(?i)^(achievements|awards|honors)",
        "languages": r"(?i)^(languages)"
    }
    
    lines = text.split('\n')
    sections = {}
    current_section = "unclassified"
    section_content = []
    
    for line in lines:
        line_clean = line.strip()
        if not line_clean:
            continue
            
        # Check if line is a potential header (short line)
        if len(line_clean) < 50:
            matched_section = None
            for sec_name, pattern in section_patterns.items():
                if re.match(pattern, line_clean):
                    matched_section = sec_name
                    break
            
            if matched_section:
                if section_content:
                    sections[current_section] = '\n'.join(section_content)
                current_section = matched_section
                section_content = []
                continue
                
        section_content.append(line_clean)
        
    if section_content:
        sections[current_section] = '\n'.join(section_content)
        
    return sections
